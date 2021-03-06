# -*- coding: utf-8 -*-

#attention:
#run command below first:
#pip install -r requirements.txt

from docx import Document
from docx.shared import Pt  #用来设置字体的大小
from docx.shared import Inches
from docx.oxml.ns import qn  #设置字体
from docx.shared import RGBColor  #设置字体的颜色
from docx.enum.style import WD_STYLE_TYPE

from docx.table import _Cell
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import json
import sys
import os
import xlrd
import re
import sys
#替换文件需要用到的包
import zipfile
import shutil
import random
import string


#解决无法导入不同目录模块的问题
sys.path.append("..")
sys.path.append(".")

from mysqlhelp import Basedb,doconfig
from returndata import returndata1



class docmaker():

    def __init__(self):
        self.name = "docmaker"
        self.tablename = "doconfig"
        self.basedb = Basedb();
        self.engine = self.basedb.crengine()
        self.session1 = self.basedb.createsession(self.engine)

        self.flagtext = '@@text\d'
        self.flagimage = '@@image\d'
        self.flagimage1 = 'image\d'
        self.flagsheet = '@@Sheet\d'    
        self.flagsheet1 = '@@excel\d'

        self.colorconfig = os.path.abspath(os.path.dirname(os.path.abspath("__file__"))) + "/src/" "color.json"
        self.colorjsondata = {}


    def set_cell_border(self,cell: _Cell, **kwargs):
        """
        Set cell`s border
        Usage:
        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            start={"sz": 24, "val": "dashed", "shadow": "true"},
            end={"sz": 12, "val": "dashed"},
        )
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
    
        # check for tag existnace, if none found, then create one
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
    
        # list over all available tags
        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)
    
                # check for tag existnace, if none found, then create one
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)
    
                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))

    def makedoc(self,doc_name):
        #read table
        docdata = self.querydocdata(doc_name)

        # print(docdata)

        print(docdata["returndt"][0].doc_template)
        print(docdata["returndt"][0].doc_outpath)
        print(docdata["returndt"][0].doc_label_text)
        print(docdata["returndt"][0].doc_image_dir)
        print(docdata["returndt"][0].doc_excel)


        #load color.json under same directory

        
        if os.path.isfile(self.colorconfig):
            load_f = open(self.colorconfig, 'r',encoding='UTF-8') 
            self.colorjsondata = json.load(load_f)

            print(self.colorjsondata)

        else:
            print(self.colorconfig  + " not exits")




        #open tempalte
        docfile =  os.path.abspath(os.path.dirname(os.path.abspath("__file__"))) + "/resouce/template/" + docdata["returndt"][0].doc_template
        print(docfile)
        document = Document(docfile)

        #read json text

        jsonfile = os.path.abspath(os.path.dirname(os.path.abspath("__file__"))) + "/resouce/text/" + docdata["returndt"][0].doc_label_text

        jsondataflag = 1

        if os.path.isfile(jsonfile):
            load_f = open(jsonfile, 'r',encoding='UTF-8') 
            jsondata = json.load(load_f)

            print(jsondata)
            jsondataflag = 0
        else:
            print(jsonfile + " not exits")
        # print(jsondata["test1"].keys())

        #get key list
        # keylist = jsondata["test1"].keys()

        # for key in keylist:
        #     print(key)
        #     print(jsondata["test1"][key])



        #read image dir 
        imagedir = os.path.abspath(os.path.dirname(os.path.abspath("__file__"))) + "/resouce/image/" + docdata["returndt"][0].doc_image_dir
        imagedirs = []
        imagefiles = []
        imagefilesnames = []
        for item in os.scandir(imagedir):
            if item.is_dir():
                imagedirs.append(item.path)

            elif item.is_file():
                imagefiles.append(item.path)
                imagefilesnames.append(item.name)

        print('\n'.join(imagefiles))
        print('\n'.join(imagefilesnames))



        #read excel
        excelfile = os.path.abspath(os.path.dirname(os.path.abspath("__file__"))) + "/resouce/excel/" + docdata["returndt"][0].doc_excel
        # excelfile = 'I://workspace1/python/docmake/resouce/excel/test1.xlsx'
        excelflag = 1
        if os.path.isfile(excelfile):

            wb = xlrd.open_workbook(excelfile)

            print("open excel "+ excelfile +" successfully")
            excelflag  = 0
        else:
            print(excelfile +" not exists")
            


        # debug
        # for sheet in wb.sheets():
        #     print(sheet.name)

        #process doc
        for paragraph in document.paragraphs:
            print("===========")
            s = paragraph.text

            searchtext = re.search(self.flagtext,s)

            searchimage = re.search(self.flagimage,s)

            searchsheet = re.search(self.flagsheet,s)
            

            #文本
            if searchtext and jsondataflag == 0 :

                print("searchtext")

                print(searchtext)

                p = re.compile(self.flagtext)

                # print(p)


                #get key list
                # keylist = jsondata["test1"].keys()
                keylist = jsondata["textdata"].keys()

                for key in keylist:
                    print("遍历")
                    print(key)
                    
                    #match text

                    searchjson = re.search(key, paragraph.text)

                    print("+++++++search start++++++++++++")
                    print(searchjson)
                    if searchjson:
                        print("got!")
                        print(key)
                        print(paragraph.text)
                    print("+++++++search end++++++++++++")

                    print(jsondata["textdata"][key])
                    print(type(jsondata["textdata"][key]))

                    # if key == paragraph.text:
                    if searchjson:

                        if isinstance(jsondata["textdata"][key],dict):
                            
                            print("this is a dict!")

                            #处理内容
                            ss = p.sub(jsondata["textdata"][key]["content"],s)
                            paragraph.text = ''
                            run = paragraph.add_run(ss)

                            #处理颜色
                            colorkey = self.colorjsondata["color"].keys() 

                            for ckey in colorkey:
                                if ckey == jsondata["textdata"][key]["color"]:
                                    run.font.color.rgb = RGBColor(self.colorjsondata["color"][ckey]['r'],self.colorjsondata["color"][ckey]['g'],self.colorjsondata["color"][ckey]['b'])
                        else:                
                            print("this is not a dict!")

                            # ss = p.sub(jsondata["test1"][key],s)
                        
                            ss = p.sub(jsondata["textdata"][key],s)
                            paragraph.text = ss
                            
                            print("matched!")
                            print(paragraph.text)
                    else:
                        print("not match")
                        pass
                
            
            else:

                print("not found text")

            #图片
            if searchimage:

                print("searchimage")

                p = re.compile(self.flagimage)

                

                for imagename in imagefilesnames:
                    
                   

                    imagenamenosuffix = imagename.split(".")

                    print(imagenamenosuffix[0])

                    searchimage1 = re.search(self.flagimage1,imagenamenosuffix[0])

                    if searchimage1:
                        print(imagenamenosuffix[0])
                        print(s[2:8])
                        if s[2:8] == imagenamenosuffix[0]:

                            

                            ss = p.sub('',s)
                            paragraph.text = ss
                            run = paragraph.add_run()
                            
                            #add picture and set image size
                            # run.add_picture(imagedir+'/'+imagename)
                            run.add_picture(imagedir+'/'+imagename, width=Inches(4))
         
            else:
    
                print("not found image")

            #电子表格
            if searchsheet and excelflag == 0:

                print("searchsheet")

                sheetname = s[2:8]

                # print(sheetname)

                for sheet in wb.sheets():
                    # print(sheet.name)

                    if sheet.name == sheetname:

                        print("debug")
                        print(sheet.name)
                        print(sheet.nrows)
                        print(sheet.ncols)

                        
                        # styles = document.styles

                        # print(styles)

                        # ss = ''
                        # for s in styles:
                        #         if s.type == WD_STYLE_TYPE.TABLE:
                        #             print(s.type)
                        #             print("++++++++")
                        #             print(s)
                        #             ss = s
                        #             print("++++++++")

                        # print("=========")
                        # print(ss)
                        
                        # print("=========")
                        # table = document.add_table(sheet.nrows,sheet.ncols,style=ss)

                        table = document.add_table(sheet.nrows,sheet.ncols)
                        

                        

                        for i in range(0,sheet.nrows):
                            for j in range(0,sheet.ncols):
                                print(i)
                                print(j)
                                print(sheet.cell_value(i,j))
                                table.cell(i,j).text = str(sheet.cell_value(i,j))

                                ##############
                                #set up cell border
                                self.set_cell_border(
                                    table.cell(i,j),
                                    top={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
                                    bottom={"sz": 12, "color": "#000000", "val": "single"},
                                    start={"sz": 24, "val": "dashed", "shadow": "true"},
                                    end={"sz": 12, "val": "dashed"},
                                )
                                ##############

                        
                        #
                        self.move_table_after(table,paragraph)
                        #
                        paragraph.text = ''
            else:
    
                print("not found sheet")


        #save doc to outpath

        savepath = os.path.abspath(os.path.dirname(os.path.abspath("__file__"))) + "/resouce/output/" + docdata["returndt"][0].doc_outpath

        print("before save")
        print(savepath)
        document.save(savepath)
        print("after save")

        attach_dir = os.path.abspath(os.path.dirname(os.path.abspath("__file__"))) + "/resouce/attachment/" + docdata["returndt"][0].doc_attach_dir
        self.subattach(savepath,attach_dir)






    #移动表格
    def move_table_after(self,table, paragraph):
        tbl, p = table._tbl, paragraph._p
        p.addnext(tbl)


    #替换word里的附件

    def subattach(self,docfile,attach_dir):

        #建立临时父目录
        if not os.path.exists(os.path.abspath(os.path.dirname(os.path.abspath("__file__"))) + "/resouce/tempdir/"):
            os.mkdir(os.path.abspath(os.path.dirname(os.path.abspath("__file__"))) + "/resouce/tempdir/")


        #以压缩格式打开word文件
        zipdoc = zipfile.ZipFile(docfile) 
        
        tmpdir = ""

        pathsuffix = os.path.abspath(os.path.dirname(os.path.abspath("__file__"))) + "/resouce/tempdir/"

        #生行8位临时文件夹名

        while True:
            tmpdir= ''.join(random.sample(string.ascii_letters + string.digits, 8))   #生行8位临时文件夹名
            if not os.path.exists(tmpdir):
                break
        tmpdir1 = pathsuffix + tmpdir

        print(tmpdir1)

        os.mkdir(tmpdir1)                  #创建临时目录
        os.chdir(tmpdir1)                   #转到临时目录
        zipdoc.extractall()                     #解压word文件到临时文件夹
        zipdoc.close()                           #关闭word文档，否则后面重新压缩会报错


        #read attach dir 
        attachdir = attach_dir
        attachdirs = []
        attachfiles = []
        attachfilesnames = []
        for item in os.scandir(attachdir):
            if item.is_dir():
                attachdirs.append(item.path)

            elif item.is_file():
                attachfiles.append(item.path)
                attachfilesnames.append(item.name)

        print('\n'.join(attachfiles))
        print('\n'.join(attachfilesnames))

        #获取docx里的嵌入对象
        embeddingdir = tmpdir1 + "/word/embeddings/"
        embeddingdirs = []
        embeddingfiles = []
        embeddingfilesnames = []
        for item in os.scandir(embeddingdir):
            if item.is_dir():
                embeddingdirs.append(item.path)

            elif item.is_file():
                embeddingfiles.append(item.path)
                embeddingfilesnames.append(item.name)

        print('\n'.join(embeddingfiles))
        print('\n'.join(embeddingfilesnames))

        #把正确文件拷贝覆盖模版文件的空附件
        i = 0
        
        for embedingfilename in embeddingfilesnames:
            j = 0
            for attachfilesname in attachfilesnames:
                if embedingfilename == attachfilesname:
                    
                    shutil.copy(attachfiles[j],embeddingfiles[i])


                j = j + 1
            
            i = i + 1

        #
        azip = zipfile.ZipFile(docfile, 'w')    #以压缩格式新建word文档
        for i in os.walk('.'):                             #使用os.walk遍历整个目录及子目录，保证原有的目录结构不变
            for j in i[2]:
                azip.write(os.path.join(i[0],j), compress_type=zipfile.ZIP_DEFLATED)     #将文件逐个打包到word文档中，压缩格式指定为ZIP_DEFLATED
        azip.close()                                       #关闭文件

        os.chdir('..')
        shutil.rmtree(tmpdir1,ignore_errors=True)    #删除临时文件夹



        

    def insertdocdata(self,doc_name,doc_template,doc_outpath,doc_label_text,doc_image_dir,doc_excel,doc_attach_dir,doc_rmrk=""):
        #check exists
        queryret = self.querydocdata(doc_name)
        if len(queryret["returndt"]) != 0:
            print("文档配置已存在")
            rd = returndata1
            rd["returncd"] = 1
            rd["returndt"] = "文档配置已存在"
            
        else:
            print("可创建文档配置")

            #check datacount
            querycount = self.querydocdatacount()
            datacount = len(querycount["returndt"])

            dc = doconfig()
            dc.doc_id = datacount + 1
            dc.doc_name = doc_name
            dc.doc_template = doc_template
            dc.doc_outpath = doc_outpath
            dc.doc_label_text = doc_label_text
            dc.doc_image_dir = doc_image_dir
            dc.doc_excel = doc_excel
            dc.doc_attach_dir = doc_attach_dir
            dc.doc_rmrk = doc_rmrk

            self.session1.add(dc)
            self.session1.commit()
            self.session1.close()

            # print(dc)
            rd = returndata1
            rd["returncd"] = 0
            rd["returndt"] = dc

        # print(rd)

        return rd



    def repairdocdata(self,doc_name,doc_template,doc_outpath,doc_label_text,doc_image_dir,doc_excel,doc_attach_dir,doc_rmrk=""):
        #check exists
        queryret = self.querydocdata(doc_name)

        if len(queryret["returndt"]) == 0:
            print("文档配置已存在")
            rd = returndata1
            rd["returncd"] = 1
            rd["returndt"] = "文档配置未找到，更新未成功"
        else:

            

            #更新数据

            dc = queryret["returndt"][0]

            dc.doc_template = doc_template
            dc.doc_outpath = doc_outpath
            dc.doc_label_text = doc_label_text
            dc.doc_image_dir = doc_image_dir
            dc.doc_excel = doc_excel
            dc.doc_attach_dir = doc_attach_dir
            dc.doc_rmrk = doc_rmrk

            self.session1.commit()
            self.session1.close()

            rd = returndata1
            rd["returncd"] = 0
            rd["returndt"] = dc

        return rd


    def deletedocdata(self,doc_name):

        dc = self.session1.query(doconfig).filter(doconfig.doc_name == doc_name).all()

        if(len(dc)!=0):

            rd = returndata1
            rd["returncd"] = 0
            rd["returndt"] = dc

            self.session1.query(doconfig).filter(doconfig.doc_name == doc_name).delete()
            self.session1.commit()
            self.session1.close()

        else:
           
            rd = returndata1
            rd["returncd"] = 1
            rd["returndt"] = "未找到数据，未删除"

        return rd


    def querydocdata(self,doc_name):
        #read table data   
        dc = self.session1.query(doconfig).filter(doconfig.doc_name == doc_name).all()

        # print("sssss")
        # print(dc)

        rd = returndata1


        rd["returncd"] = 0
        rd["returndt"] = dc


        # print(dc.doc_template)    
        return rd


    def querydocdatacount(self):
        dc = self.session1.query(doconfig)
        # print(len(dc.all()))
        # return dc.all()
        rd = returndata1
        rd["returncd"] = 0
        rd["returndt"] = dc.all()

        print(dc.all())

        return rd
        
           

if __name__ == '__main__':

    # print("start python script")

    para = "test1"
    # para = "test5"
    
    # print(sys.argv)
    if len(sys.argv) > 1:

        # print("进程 " +sys.argv[1] +" 执行。") 
        para = sys.argv[1]
    else:

        # print("无参数")
        pass



    dm = docmaker()
    # dm.querydocdata("test1")
    # dm.querydocdatacount();
    # ret = dm.insertdocdata("test4","template","fasdfs","ddddd","doc_image_dir","doc_excel","doc_rmrk")

    # ret = dm.repairdocdata("test4",para,"fasdfs","ddddd","doc_image_dir","doc_excel","doc_rmrk")

    # if(ret["returncd"]) == 0:

    #     # print(ret["returndt"].doc_id)

    #     # print(ret)
    #     # print(ret["returndt"].doc_id)
    #     # print(ret["returndt"].doc_name)
    #     # print(ret["returndt"].doc_template)
    #     # print(ret["returndt"].doc_outpath)
    #     # print(ret["returndt"].doc_label_text)
    #     # print(ret["returndt"].doc_image_dir)
    #     # print(ret["returndt"].doc_excel)
    #     # print(ret["returndt"].doc_rmrk)

    #     print(ret["returndt"])
    
    # else:

    #     print("未插入数据")


    dm.makedoc(para)

    # dm.querydocdatacount();

    



    
